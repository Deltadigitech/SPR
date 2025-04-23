from flask import Flask, render_template, request, jsonify, session
from flask_sqlalchemy import SQLAlchemy
import os
import io
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from config import Config
import openai
import re
from datetime import timedelta

# Initialize Flask app and configure the database
app = Flask(__name__)
app.config.from_object(Config)
app.secret_key = 'your_secret_key_here'
app.permanent_session_lifetime = timedelta(minutes=30)

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

# Set OpenAI API key
openai.api_key = app.config['OPENAI_API_KEY']

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    email = db.Column(db.String(100), nullable=True)
    phone = db.Column(db.String(20), nullable=True)

class ChatHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_message = db.Column(db.String(255), nullable=False)
    bot_response = db.Column(db.String(255), nullable=False)
    timestamp = db.Column(db.DateTime, default=db.func.current_timestamp())

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    file_data = db.Column(db.LargeBinary, nullable=False)
    text_content = db.Column(db.Text, nullable=True)
    uploaded_at = db.Column(db.DateTime, default=db.func.current_timestamp())

# Helper function to extract text from documents
def extract_text_from_file(file, filename):
    if filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(file))
        text = ''
        for page_num in range(pdf_reader.numPages):
            page = pdf_reader.getPage(page_num)
            text += page.extract_text() if page.extract_text() else ''
        return text
    elif filename.endswith('.docx'):
        doc = docx.Document(io.BytesIO(file))
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        return None

# Helper function to search for answers in document content
def search_document_content(query):
    documents = Document.query.all()
    if not documents:
        return None

    document_texts = [doc.text_content for doc in documents if doc.text_content]
    if not document_texts:
        return None

    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(document_texts + [query])
    cosine_similarities = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])
    most_relevant_doc_index = cosine_similarities.argmax()

    return document_texts[most_relevant_doc_index]

def replace_markdown_bold(text):
    pattern = r"\*\*(.*?)\*\*"
    return re.sub(pattern, r"<b>\1</b>", text)

def format_gpt_response(raw_response):
    text = replace_markdown_bold(raw_response)
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    if len(lines) > 1:
        formatted_lines = []
        for line in lines:
            if not (line.startswith('- ') or line.startswith('* ')):
                formatted_lines.append(f"- {line}")
            else:
                formatted_lines.append(line)
        text = "\n".join(formatted_lines)
    else:
        text = lines[0] if lines else ""

    return text

# Function to generate a response using OpenAI's GPT model
def generate_gpt_response(query, context, detailed=False):
    try:
        if detailed:
            prompt = f"Context: {context}\n\nQuestion: {query}\n\nExplain in detail:"
        else:
            prompt = f"Context: {context}\n\nQuestion: {query}\n\nProvide a short and concise answer:"

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini-2024-07-18",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a friendly and knowledgeable AI assistant for Delta Media. "
                        "You represent Delta Media and should always respond from the first-person perspective, "
                        "using words like 'we,' 'our,' and 'us.'\n\n"
                        "Rules:\n"
                        "1. Always respond as if you are Delta Media. Use 'we,' 'our,' and 'us' instead of 'they' or 'their.'\n"
                        "2. Don't justify your answers.\n"
                        "3. Don't give information not mentioned in the CONTEXT INFORMATION.\n"
                        "4. You can ask counter questions if you are not clear or in doubt.\n"
                        "5. Maintain context within a conversation. If a user inquires about a specific service, "
                        "remember that question as the current context for subsequent questions.\n"
                        "6. Avoid speculation: Avoid adding details from outside the context.\n"
                        "7. If unrelated questions are asked, politely refuse to answer and bring the conversation back to Delta Media.\n\n"
                        "Format your responses as follows:\n"
                        "- Use **bold** for emphasis.\n"
                        "- Keep responses short and crisp.\n"
                        "- Use hyphens only when listing multiple points.\n"
                        "- Provide a brief definition or explanation for each topic when relevant."
                    )
                },
                {"role": "user", "content": prompt}
            ]
        )

        raw_text = response['choices'][0]['message']['content']
        return format_gpt_response(raw_text)
    except Exception as e:
        return f"An error occurred while generating the response: {str(e)}"

#Home Route
@app.route('/')
def index():
    return render_template('index.html')

# Chatbot Route
@app.route('/chatbot')
def chatbot():
    # Initialize session for new users
    if 'conversation_context' not in session:
        session['conversation_context'] = ""
    if 'user_info_asked' not in session:
        session['user_info_asked'] = False
    if 'user_info_provided' not in session:
        session['user_info_provided'] = False
    if 'chat_count' not in session:
        session['chat_count'] = 0
    return render_template('chatbot.html')

# Chat Route
@app.route('/chat', methods=['POST'])
def chat():
    user_message = request.form['message'].lower()

    # Increment chat count
    session['chat_count'] += 1
    session.modified = True

    # Check for greetings
    greetings = ['hi', 'hello', 'hey', 'hi there', 'hello there']
    if user_message in greetings:
        found_answer = "Hello! How can we assist you today?"
    else:
        # Search for relevant content in documents
        context = search_document_content(user_message)
        if context:
            # Check if the user is asking for a detailed explanation
            detailed = "explain in detail" in user_message or "elaborate" in user_message
            found_answer = generate_gpt_response(user_message, context, detailed)
        else:
            found_answer = "We couldn't find any relevant information in our documents."

    # Store chat history in the database
    chat_entry = ChatHistory(user_message=user_message, bot_response=found_answer)
    db.session.add(chat_entry)
    db.session.commit()

    # Update conversation context in session
    session['conversation_context'] += f"User: {user_message}\nBot: {found_answer}\n"
    session.modified = True

    # Ask for user info only if not already provided and after 2-3 interactions
    ask_user_info = (
        not session['user_info_provided']  # User details not provided yet
        and session['chat_count'] >= 3  # After 2-3 interactions
        and not session['user_info_asked']  # Not already asked in this session
    )

    if ask_user_info:
        session['user_info_asked'] = True
        found_answer += "\n\nCould you please provide your name, email, and phone number so we can assist you better?"

    return jsonify({
        'bot_response': found_answer,
        'ask_user_info': ask_user_info
    })

# Store User Info
@app.route('/store_user_info', methods=['POST'])
def store_user_info():
    name = request.form['name']
    email = request.form['email']
    phone = request.form['phone']

    user = User(name=name, email=email, phone=phone)
    db.session.add(user)
    db.session.commit()

    # Mark user info as provided in the session
    session['user_info_provided'] = True
    session.modified = True

    return jsonify({'bot_response': "Thanks for your info!"})

# Upload Document Route
@app.route('/upload', methods=['GET', 'POST'])
def upload_page():
    if request.method == 'POST':
        if 'document' not in request.files:
            return "No file part", 400
        file = request.files['document']
        if file.filename == '':
            return "No selected file", 400

        filename = secure_filename(file.filename)
        file_data = file.read()

        text_content = extract_text_from_file(file_data, filename)
        if not text_content:
            return "Unsupported file type", 400

        new_document = Document(filename=filename, file_data=file_data, text_content=text_content)
        db.session.add(new_document)
        db.session.commit()

        return render_template('upload.html', message="✅ File uploaded successfully!")

    return render_template('upload.html')

# Chat History Route
@app.route('/chat_history')
def chat_history():
    chat_entries = ChatHistory.query.order_by(ChatHistory.timestamp.desc()).all()
    return render_template('chat_history.html', chat_entries=chat_entries)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)